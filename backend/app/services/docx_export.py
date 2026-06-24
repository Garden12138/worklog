from io import BytesIO

from docx import Document


def markdown_to_docx(markdown: str) -> BytesIO:
    document = Document()
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
            continue
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            document.add_heading(text, min(max(level, 1), 4))
            continue
        if stripped.startswith(("- ", "* ")):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        if len(stripped) > 3 and stripped[0].isdigit() and ". " in stripped[:5]:
            document.add_paragraph(stripped.split(". ", 1)[1].strip(), style="List Number")
            continue
        document.add_paragraph(stripped)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
